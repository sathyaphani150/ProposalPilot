from __future__ import annotations

from io import BytesIO
from typing import Any


def build_docx_bytes(proposal: dict[str, Any]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Proposal", level=0)
    sections = [
        ("Executive Summary", proposal.get("executive_summary")),
        ("Problem Statement", proposal.get("client_problem_statement")),
        ("Solution", proposal.get("proposed_solution")),
        ("Architecture", proposal.get("technical_architecture")),
        ("Technology Stack", proposal.get("technology_stack")),
        ("Delivery Approach", proposal.get("delivery_approach")),
        ("Cost Estimate", proposal.get("cost_estimation")),
        ("Risks", "\n".join(proposal.get("risks") or [])),
        ("Assumptions", "\n".join(proposal.get("assumptions") or [])),
    ]
    for title, value in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(str(value or ""))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

