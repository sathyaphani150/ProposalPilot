"""
ProposalPilot AI — Document Parser Service
Extracts clean text from PDF, DOCX, TXT, and Markdown files.
Normalises whitespace and returns structured text suitable for LLM processing.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from loguru import logger

from app.exceptions import DocumentParsingError, InvalidFileTypeError
from app.config import get_settings

settings = get_settings()

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


async def parse_document(file_path: str | Path) -> str:
    """
    Parse a document and return clean text content.

    Args:
        file_path: Absolute path to the uploaded file.

    Returns:
        Clean, normalised text content.

    Raises:
        InvalidFileTypeError: If the file extension is not supported.
        DocumentParsingError: If parsing fails for any reason.
    """
    path = Path(file_path)

    if not path.exists():
        raise DocumentParsingError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise InvalidFileTypeError(
            f"File type '{ext}' is not supported. Allowed: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    logger.debug(f"Parsing document: {path.name} ({ext})")

    try:
        if ext == ".pdf":
            text = _parse_pdf(path)
        elif ext == ".docx":
            text = _parse_docx(path)
        elif ext in (".txt", ".md"):
            text = _parse_text(path)
        else:
            raise DocumentParsingError(f"No parser available for {ext}")
    except (InvalidFileTypeError, DocumentParsingError):
        raise
    except Exception as e:
        logger.error(f"Unexpected error parsing {path.name}: {e}")
        raise DocumentParsingError(
            f"Failed to parse '{path.name}': {str(e)}"
        )

    cleaned = _clean_text(text)

    if not cleaned.strip():
        raise DocumentParsingError(
            f"No readable text could be extracted from '{path.name}'. "
            "The file may be scanned, image-based, or password-protected."
        )

    logger.info(f"Parsed '{path.name}': {len(cleaned)} characters extracted")
    return cleaned


def _parse_pdf(path: Path) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise DocumentParsingError(
            "PyMuPDF is not installed. Run: pip install pymupdf"
        )

    doc: Any = fitz.open(str(path))
    pages: list[str] = []

    for page_index in range(int(doc.page_count)):
        page = doc.load_page(page_index)
        page_text = page.get_text("text")  # type: ignore[attr-defined]
        if page_text.strip():
            pages.append(f"--- Page {page_index + 1} ---\n{page_text}")

    doc.close()
    return "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise DocumentParsingError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    doc = Document(str(path))
    sections: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            sections.append(para.text.strip())

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                sections.append(row_text)

    return "\n\n".join(sections)


def _parse_text(path: Path) -> str:
    """Read plain text or Markdown file."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParsingError(
        f"Could not decode '{path.name}' with any known encoding."
    )


def _clean_text(text: str) -> str:
    """
    Normalise whitespace, remove null bytes, and strip excessive blank lines.
    """
    import re

    # Remove null bytes
    text = text.replace("\x00", "")
    # Normalise line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse more than 3 consecutive newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove trailing whitespace per line
    lines = [line.rstrip() for line in text.split("\n")]
    return "\n".join(lines).strip()


def get_file_extension(filename: str) -> str:
    """Return lowercase extension without dot."""
    return Path(filename).suffix.lower().lstrip(".")


def validate_upload_file(filename: str, file_size: int) -> None:
    """
    Validate file type and size before saving.

    Raises:
        InvalidFileTypeError: If extension not in allowlist.
        FileTooLargeError: If file exceeds max size.
    """
    from app.exceptions import FileTooLargeError

    ext = get_file_extension(filename)
    if ext not in [e.lstrip(".") for e in SUPPORTED_EXTENSIONS]:
        raise InvalidFileTypeError(
            f"File type '.{ext}' not supported. Allowed: pdf, docx, txt, md"
        )

    if file_size > settings.max_upload_size_bytes:
        raise FileTooLargeError(
            f"File size {file_size / 1024 / 1024:.1f}MB exceeds "
            f"maximum {settings.MAX_UPLOAD_SIZE_MB}MB."
        )
